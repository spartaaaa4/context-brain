import os
import io
import json
import base64
import threading
import re
from anthropic import Anthropic
from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception
import hashlib
from datetime import datetime
from models import db, Message, Document, Vertical, Note, ProcessMap, ProcessMapFeedback, VerticalIntelligence, IntelligenceFeedback, IntelligenceSection

ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY")

client = Anthropic(
    api_key=ANTHROPIC_API_KEY,
)

# In-memory tracking for async intelligence generation
_intel_generation_status = {}


def get_intel_generation_status(vertical_id):
    return _intel_generation_status.get(vertical_id, {})


def clean_json_response(text):
    """Strip markdown code fences and extract raw JSON from Claude responses."""
    if not text:
        return text
    cleaned = text.strip()
    if cleaned.startswith("```"):
        first_newline = cleaned.find("\n")
        if first_newline != -1:
            cleaned = cleaned[first_newline + 1:]
        if cleaned.endswith("```"):
            cleaned = cleaned[:-3].strip()
    return cleaned

DISCOVERY_PHASES = [
    {
        "phase": "Business Model & Scale",
        "keywords": [
            "revenue", "pricing", "client", "customer", "margin", "volume",
            "ticket", "deal", "payment", "scale", "business model"
        ],
    },
    {
        "phase": "Supply Side",
        "keywords": [
            "worker", "candidate", "sourcing", "recruitment", "hiring",
            "screening", "onboarding", "qualification", "pool"
        ],
    },
    {
        "phase": "Matching & Deployment",
        "keywords": [
            "matching", "allocation", "deployment", "assignment", "shift",
            "dispatch", "placement", "fulfillment", "schedule"
        ],
    },
    {
        "phase": "Money Flow",
        "keywords": [
            "payroll", "payout", "billing", "invoice", "payment", "salary",
            "commission", "reconciliation", "settlement"
        ],
    },
    {
        "phase": "Team & Operations",
        "keywords": [
            "team", "ops", "manager", "qc", "quality", "review", "admin",
            "headcount", "manual", "compliance", "support"
        ],
    },
]


def safe_json_loads(raw_value, default=None):
    if raw_value is None:
        return {} if default is None else default
    if isinstance(raw_value, (dict, list)):
        return raw_value
    try:
        return json.loads(raw_value)
    except (TypeError, json.JSONDecodeError):
        return {} if default is None else default


def get_currency_context(vertical):
    geography = (vertical.geography or "").lower()
    if "malaysia" in geography or "singapore" in geography:
        return {"symbol": "RM", "monthly_salary": 3500, "hourly_salary": 3500 / 176, "salary_text": "~RM3,500/month"}
    if "indonesia" in geography:
        return {"symbol": "Rp", "monthly_salary": 5000000, "hourly_salary": 5000000 / 176, "salary_text": "~Rp5,000,000/month"}
    return {"symbol": "₹", "monthly_salary": 25000, "hourly_salary": 25000 / 176, "salary_text": "~₹25,000/month"}


def build_validated_facts(vertical_id):
    facts = []
    feedback_entries = IntelligenceFeedback.query.filter_by(vertical_id=vertical_id).order_by(IntelligenceFeedback.created_at.desc()).all()
    for fb in feedback_entries:
        if fb.feedback_type in {"correct", "edit"}:
            field_label = fb.field_path or fb.section or "general"
            if fb.corrected_value:
                facts.append(f"{field_label}: {fb.corrected_value}")
            elif fb.original_value:
                facts.append(f"{field_label}: {fb.original_value}")
        elif fb.feedback_type == "partially_correct" and fb.corrected_value:
            facts.append(f"{fb.field_path or fb.section}: {fb.corrected_value}")

    notes = Note.query.filter_by(vertical_id=vertical_id).order_by(Note.created_at.desc()).all()
    for note in notes:
        if note.content.startswith("[Knowledge Gap Answer]"):
            facts.append(note.content.replace("[Knowledge Gap Answer] ", "").strip())

    deduped = []
    seen = set()
    for fact in facts:
        normalized = fact.strip().lower()
        if normalized and normalized not in seen:
            seen.add(normalized)
            deduped.append(fact.strip())
    return deduped[:12]


def compute_discovery_coverage(vertical_id):
    messages = Message.query.filter_by(vertical_id=vertical_id).all()
    all_text = " ".join((m.content or "").lower() for m in messages)
    coverage = []
    for phase in DISCOVERY_PHASES:
        keywords = phase["keywords"]
        hits = sum(1 for keyword in keywords if keyword in all_text)
        total = len(keywords) or 1
        percentage = min(100, int((hits / total) * 130))
        if percentage > 60:
            status = "well_covered"
            detail = "well covered"
        elif percentage > 25:
            status = "partial"
            detail = "partially covered"
        else:
            status = "not_covered"
            detail = "not yet discussed"
        coverage.append({
            "phase": phase["phase"],
            "percentage": percentage,
            "status": status,
            "detail": detail,
        })
    return coverage


def build_chat_intelligence_context(vertical_id):
    latest_intel = VerticalIntelligence.query.filter_by(vertical_id=vertical_id).order_by(VerticalIntelligence.generated_at.desc()).first()
    intelligence = safe_json_loads(latest_intel.intelligence_data if latest_intel else None, default={})
    confirmed_facts = build_validated_facts(vertical_id)
    coverage = compute_discovery_coverage(vertical_id)
    least_covered = [item["phase"] for item in coverage if item["status"] != "well_covered"][:2]
    knowledge_gaps = intelligence.get("knowledgeGaps") or []

    sections = []

    if confirmed_facts:
        sections.append("CONFIRMED FACTS (do not re-ask these):")
        sections.extend(f"- {fact}" for fact in confirmed_facts[:8])

    if knowledge_gaps:
        sections.append("")
        sections.append("VALIDATION / OPEN QUESTIONS:")
        for gap in knowledge_gaps[:5]:
            if isinstance(gap, dict):
                question = gap.get("question") or gap.get("content") or str(gap)
            else:
                question = str(gap)
            sections.append(f"- {question}")

    if coverage:
        sections.append("")
        sections.append("DISCOVERY COVERAGE:")
        status_icon = {"well_covered": "✅", "partial": "⚠️", "not_covered": "❌"}
        for item in coverage:
            sections.append(
                f"{status_icon[item['status']]} {item['phase']} — {item['detail']} ({item['percentage']}%)"
            )

    if least_covered:
        sections.append("")
        sections.append(f"PRIORITY: Focus the next few questions on {', '.join(least_covered)}.")

    return "\n".join(sections).strip()


def is_rate_limit_error(exception):
    error_msg = str(exception)
    return (
        "429" in error_msg
        or "RATELIMIT_EXCEEDED" in error_msg
        or "quota" in error_msg.lower()
        or "rate limit" in error_msg.lower()
        or (hasattr(exception, "status_code") and exception.status_code == 429)
    )


def get_chat_system_prompt(vertical, intelligence_context=""):
    return f"""You are an AI process analyst working for BetterPlace Group's Central AI Labs team. You are conducting a structured intake to understand the operations of {vertical.name} ({vertical.geography} — {vertical.type}).

{vertical.seed_context}

Your goal is to build a complete picture of this organization's operations. You need to understand:

1. BUSINESS MODEL — What exactly does this business do? How does it make money? Who are the customers? What is the scale?
2. TEAM & ORG STRUCTURE — Who works here? What are the key roles? How many people? Who reports to whom?
3. END-TO-END PROCESSES — Walk through every major workflow step by step. Especially recruitment/staffing flows.
4. TOOLS & SYSTEMS — What software, apps, spreadsheets, WhatsApp groups, manual processes do they use?
5. BOTTLENECKS & PAIN POINTS — Where is time wasted? Where do errors happen? What's frustrating?
6. VOLUME & SCALE — How many workers/tasks/shifts/hires per day/week/month? Seasonality?
7. COMMUNICATION CHANNELS — How do they reach workers? Calls, WhatsApp, app notifications, SMS?
8. QUALITY & COMPLIANCE — How is work quality verified? What compliance requirements exist?
9. COSTS — Where is money spent on operations? What's the team cost? Tool costs?

INTERVIEW STYLE:
- Be warm, professional, and conversational — never robotic or clinical
- Ask ONE focused question at a time
- After each answer, briefly acknowledge what you heard (1 sentence), then ask the natural next question
- When something interesting is mentioned, dig deeper with a follow-up before moving to a new topic
- Every 5-6 exchanges, provide a brief summary of what you have captured so far
- Keep your responses concise — 2-3 sentences before your question, never more
- If someone provides vague information, ask for a specific example or a walk-through of a real scenario
- Use their terminology — don't correct their language
- If someone provides irrelevant or off-topic information, gently redirect: acknowledge what they said, then steer back with "That's helpful context. Going back to [topic] — could you tell me about [specific question]?"
- If someone gives very short answers, encourage them: "Could you walk me through a real example of how that works day-to-day?"

STRUCTURED QUESTIONS (MCQ / Multi-select):
When a question has a natural set of possible answers (e.g., categories, frequency ranges, yes/no, tools from a known list), present it as a structured clickable question using this exact format:

For single-select (user picks one):
[MCQ]
What is the average time to fill a position?
- Less than 24 hours
- 1-3 days
- 4-7 days
- More than a week
[/MCQ]

For multi-select (user picks multiple):
[MCQ:multi]
Which communication channels does your team use?
- WhatsApp
- Phone calls
- Email
- SMS
- In-app notifications
- Other
[/MCQ]

Rules for structured questions:
- Use MCQ when the answer is likely one of several known categories or ranges
- Use MCQ:multi when the user might select more than one option
- Always include an open-ended option like "Other" when the list might not be exhaustive
- Mix structured questions with open-ended ones — don't make every question an MCQ
- After the user selects an option, you can follow up with an open-ended question to dig deeper
- You can include regular text before or after an MCQ block
- Use these especially for: scale/volume ranges, frequency, tool lists, yes/no with nuance, severity ratings, categories

START by greeting them warmly, acknowledging their organization, and asking them to describe what their business actually does in their own words — even if the seed context already describes it. Their perspective matters.

CRITICAL: You are gathering intelligence to help build AI automation. Pay extra attention to manual, repetitive, time-consuming activities — those are the highest-value automation targets. When you hear something manual, always ask: how long does it take, how often, and how many people are involved.

{intelligence_context or ""}
"""


@retry(
    stop=stop_after_attempt(3),
    wait=wait_exponential(multiplier=1, min=2, max=30),
    retry=retry_if_exception(is_rate_limit_error),
    reraise=True
)
def send_chat_message(vertical, messages_history, user_message):
    intelligence_context = build_chat_intelligence_context(vertical.id)
    system_prompt = get_chat_system_prompt(vertical, intelligence_context)

    api_messages = []
    for msg in messages_history:
        api_messages.append({
            "role": msg.role,
            "content": msg.content
        })

    if len(api_messages) > 50:
        summary_msgs = api_messages[:-30]
        recent_msgs = api_messages[-30:]
        summary_text = "\n".join([f"{m['role']}: {m['content'][:200]}" for m in summary_msgs])
        api_messages = [{"role": "user", "content": f"[Summary of earlier conversation:\n{summary_text}\n]"},
                        {"role": "assistant", "content": "Thank you for the context. Let me continue our conversation."}] + recent_msgs

    api_messages.append({"role": "user", "content": user_message})

    response = client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=1024,
        system=system_prompt,
        messages=api_messages
    )

    return response.content[0].text


def get_document_extraction_prompt(vertical, doc_type, user_description):
    return f"""You are a senior analyst extracting deep intelligence from a document uploaded by a team member from {vertical.name} ({vertical.geography} — {vertical.type}).

The user described this document as: "{user_description}"
Document type: {doc_type}

Extract EVERY piece of useful information. Be exhaustive. Preserve specific numbers, names, and quotes exactly as they appear. Return ONLY valid JSON (no markdown, no backticks):

{{
  "summary": "Comprehensive 5-10 sentence summary covering the document's purpose, scope, and key findings",
  "processFlows": [
    {{
      "processName": "Name of this process or workflow",
      "steps": [
        {{
          "name": "Step name",
          "description": "Detailed description of what happens",
          "owner": "Who is responsible (role or team name) or null",
          "tools": ["Tools/systems used in this step"],
          "preconditions": "What must be true before this step or null",
          "outputs": "What this step produces or null",
          "decisionPoints": "Any branching logic or conditions or null",
          "timeEstimate": "How long this step takes or null",
          "frequency": "How often this happens (daily, per request, etc.) or null",
          "volume": "Number of times per day/week/month or null",
          "exceptions": "What happens when things go wrong or null",
          "dependencies": ["Other steps or systems this depends on"]
        }}
      ]
    }}
  ],
  "organizationalInfo": {{
    "teamStructure": [
      {{
        "role": "Role or title",
        "headcount": "Number of people or null",
        "responsibilities": "Key responsibilities",
        "reportsTo": "Who they report to or null",
        "processSteps": ["Which process steps they own"]
      }}
    ],
    "reportingLines": "Description of org hierarchy if mentioned or null",
    "totalHeadcount": "Total team size if mentioned or null"
  }},
  "financialData": {{
    "costs": [
      {{
        "item": "What the cost is for",
        "amount": "Exact amount as stated in document",
        "period": "per month, per year, per transaction, etc. or null",
        "context": "Additional context about this cost"
      }}
    ],
    "pricing": [
      {{
        "model": "Pricing model description",
        "amount": "Price point or range",
        "context": "Who pays, when, for what"
      }}
    ],
    "budgets": "Any budget allocations mentioned or null",
    "unitEconomics": "Revenue or cost per unit/transaction if mentioned or null"
  }},
  "clientInfo": [
    {{
      "name": "Client or customer name",
      "industry": "Industry or segment or null",
      "size": "Company size or volume or null",
      "sla": "Service level requirements or null",
      "relationship": "Nature of the relationship or null"
    }}
  ],
  "technologyStack": [
    {{
      "name": "Tool, app, or system name",
      "purpose": "What it is used for",
      "usedBy": "Who uses it or null",
      "integrations": ["Other systems it connects to"],
      "limitations": "Known limitations or complaints or null"
    }}
  ],
  "painPointsAndChallenges": [
    {{
      "issue": "Description of the pain point or challenge",
      "severity": "high|medium|low — based on language and context",
      "affectedArea": "Which process or team is affected",
      "currentWorkaround": "How they handle it now or null",
      "impact": "Business impact (time wasted, cost, errors, etc.) or null"
    }}
  ],
  "metricsAndKPIs": [
    {{
      "metric": "Name of the metric",
      "value": "The exact value as stated",
      "period": "Time period or null",
      "context": "What this metric measures and why it matters"
    }}
  ],
  "complianceAndRegulatory": [
    {{
      "requirement": "Regulatory or compliance requirement",
      "applicability": "When or to whom it applies",
      "procedure": "How they comply or null",
      "risk": "What happens if not complied with or null"
    }}
  ],
  "directQuotes": [
    {{
      "quote": "Exact verbatim text from the document",
      "context": "Why this quote is significant",
      "speaker": "Who said it, if identifiable, or null"
    }}
  ],
  "specificNumbers": [
    {{
      "value": "The exact number or statistic",
      "meaning": "What this number represents",
      "source": "Where in the document this appeared"
    }}
  ],
  "relationships": [
    {{
      "from": "Entity A (person, team, system, process)",
      "to": "Entity B",
      "type": "Nature of relationship (reports to, feeds into, integrates with, depends on, etc.)"
    }}
  ],
  "keyFacts": ["Other important facts not captured above — preserve detail"],
  "relevanceScore": "high|medium|low"
}}

If this is a meeting transcript, also include:
- "decisions": [{{"decision": "What was decided", "owner": "Who is responsible", "deadline": "When, or null"}}]
- "actionItems": [{{"item": "Action to take", "owner": "Who", "deadline": "When, or null", "status": "open"}}]
- "discussionTopics": [{{"topic": "Subject discussed", "keyPoints": ["Main points made"], "outcome": "Resolution or next step"}}]

CRITICAL RULES:
- Preserve ALL specific numbers exactly (amounts, percentages, dates, volumes, headcounts)
- Preserve ALL names (people, companies, tools, systems, locations)
- For any field where the document provides no information, use null — do NOT guess
- Include direct quotes for any particularly insightful or important statements
- Extract relationships between entities whenever possible
- Be exhaustive — it is better to extract too much than too little"""


@retry(
    stop=stop_after_attempt(3),
    wait=wait_exponential(multiplier=1, min=2, max=30),
    retry=retry_if_exception(is_rate_limit_error),
    reraise=True
)
def process_document_content(doc, vertical, text_content=None, base64_content=None, media_type=None):
    system_prompt = get_document_extraction_prompt(vertical, doc.doc_type, doc.user_description or "")

    messages_content = []
    if base64_content and media_type:
        if media_type.startswith("image/"):
            messages_content.append({
                "type": "image",
                "source": {
                    "type": "base64",
                    "media_type": media_type,
                    "data": base64_content,
                }
            })
        else:
            messages_content.append({
                "type": "document",
                "source": {
                    "type": "base64",
                    "media_type": media_type,
                    "data": base64_content,
                }
            })
        messages_content.append({
            "type": "text",
            "text": "Please analyze this document and extract structured information."
        })
    elif text_content:
        messages_content.append({
            "type": "text",
            "text": f"Document content:\n\n{text_content}\n\nPlease analyze this document and extract structured information."
        })
    else:
        return None

    response = client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=4096,
        system=system_prompt,
        messages=[{"role": "user", "content": messages_content}]
    )

    return response.content[0].text


def extract_pdf_text(doc):
    """Extract text from a PDF using PyPDF2 as fallback for large files."""
    try:
        from PyPDF2 import PdfReader
        if doc.file_data:
            reader = PdfReader(io.BytesIO(base64.b64decode(doc.file_data)))
        elif doc.file_path:
            reader = PdfReader(doc.file_path)
        else:
            return None
        text_parts = []
        for i, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_parts.append(f"--- Page {i + 1} ---\n{page_text}")
            except Exception:
                continue
            if len("\n".join(text_parts)) > 80000:
                break
        full_text = "\n".join(text_parts)
        return full_text[:80000] if full_text.strip() else None
    except Exception as e:
        print(f"PyPDF2 extraction failed for {doc.filename}: {e}")
        return None


MAX_PDF_SIZE_FOR_VISION = 10 * 1024 * 1024  # 10 MB — above this, use text extraction


def process_document_background(app, doc_id):
    with app.app_context():
        doc = Document.query.get(doc_id)
        if not doc:
            return

        doc.processing_status = 'processing'
        db.session.commit()

        try:
            vertical = Vertical.query.get(doc.vertical_id)
            file_path = doc.file_path

            text_content = None
            base64_content = None
            media_type = None

            if doc.file_type in ('txt', 'csv'):
                if doc.file_data:
                    text_content = base64.b64decode(doc.file_data).decode('utf-8', errors='ignore')[:80000]
                else:
                    with open(file_path, 'r', errors='ignore') as f:
                        text_content = f.read()[:80000]
            elif doc.file_type == 'xlsx':
                import openpyxl
                if doc.file_data:
                    wb = openpyxl.load_workbook(io.BytesIO(base64.b64decode(doc.file_data)))
                else:
                    wb = openpyxl.load_workbook(file_path)
                text_parts = []
                for sheet in wb.worksheets:
                    text_parts.append(f"=== Sheet: {sheet.title} ===")
                    rows = list(sheet.iter_rows(values_only=True))
                    if rows:
                        headers = [str(c) if c is not None else "" for c in rows[0]]
                        text_parts.append("Headers: " + " | ".join(headers))
                        for row in rows[1:]:
                            row_data = []
                            for i, cell in enumerate(row):
                                header = headers[i] if i < len(headers) else f"Col{i}"
                                val = str(cell) if cell is not None else ""
                                if val:
                                    row_data.append(f"{header}: {val}")
                            if row_data:
                                text_parts.append(", ".join(row_data))
                text_content = "\n".join(text_parts)[:80000]
            elif doc.file_type == 'docx':
                import docx
                if doc.file_data:
                    document = docx.Document(io.BytesIO(base64.b64decode(doc.file_data)))
                else:
                    document = docx.Document(file_path)
                text_parts = []
                for p in document.paragraphs:
                    if p.text.strip():
                        text_parts.append(p.text)
                for table in document.tables:
                    text_parts.append("\n[Table]")
                    for row in table.rows:
                        row_text = " | ".join([cell.text.strip() for cell in row.cells])
                        text_parts.append(row_text)
                text_content = "\n".join(text_parts)[:80000]
            elif doc.file_type == 'pdf':
                raw_size = doc.file_size or 0
                if raw_size > MAX_PDF_SIZE_FOR_VISION:
                    # Large PDF — extract text with PyPDF2 instead of sending binary
                    print(f"PDF {doc.filename} is {raw_size / 1024 / 1024:.1f}MB — using text extraction")
                    text_content = extract_pdf_text(doc)
                    if not text_content:
                        # If text extraction yielded nothing, try vision anyway (will likely fail)
                        print(f"Text extraction empty for {doc.filename}, trying vision API")
                        if doc.file_data:
                            base64_content = doc.file_data
                        elif file_path:
                            with open(file_path, 'rb') as f:
                                base64_content = base64.b64encode(f.read()).decode('utf-8')
                        media_type = "application/pdf"
                else:
                    if doc.file_data:
                        base64_content = doc.file_data
                    elif file_path:
                        with open(file_path, 'rb') as f:
                            base64_content = base64.b64encode(f.read()).decode('utf-8')
                    media_type = "application/pdf"
            elif doc.file_type in ('image', 'png', 'jpg', 'jpeg'):
                ext_map = {'png': 'image/png', 'jpg': 'image/jpeg', 'jpeg': 'image/jpeg', 'image': 'image/jpeg'}
                media_type = ext_map.get(doc.file_type, 'image/jpeg')
                if doc.file_data:
                    base64_content = doc.file_data
                else:
                    with open(file_path, 'rb') as f:
                        base64_content = base64.b64encode(f.read()).decode('utf-8')

            # Store raw full text for intelligence context
            if text_content:
                doc.full_text = text_content
            elif doc.file_type == 'pdf' and base64_content:
                # For small PDFs sent via vision API, also extract text as backup
                pdf_text = extract_pdf_text(doc)
                if pdf_text:
                    doc.full_text = pdf_text
            db.session.commit()

            result = process_document_content(doc, vertical, text_content, base64_content, media_type)
            if result:
                doc.extracted_content = result
                doc.processing_status = 'done'
                db.session.commit()
                start_incremental_intelligence(app, doc.vertical_id, user_id=doc.user_id)
            else:
                doc.processing_status = 'failed'
                db.session.commit()
        except Exception as e:
            print(f"Document processing error for {doc.filename}: {e}")
            doc.processing_status = 'failed'
            db.session.commit()


def start_document_processing(app, doc_id):
    thread = threading.Thread(target=process_document_background, args=(app, doc_id))
    thread.daemon = True
    thread.start()


def get_process_map_prompt(vertical):
    # Single Claude API call returns all 6 sections: businessOverview, processMap, keyInsights, topAutomationTargets, knowledgeGaps, communicationChannels/complianceNotes
    return f"""You are a senior process analyst at BetterPlace Group's AI Labs. You have been given all available context about {vertical.name} ({vertical.geography} — {vertical.type}).

Your job is to produce a comprehensive, structured analysis of this business unit.

Produce a complete analysis. Return ONLY valid JSON (no markdown, no backticks):

{{
  "businessOverview": {{
    "summary": "2-3 paragraph description of what this business does and how it operates",
    "businessModel": "How they make money",
    "scale": {{
      "workers": "Size of worker network",
      "clients": "Number and type of clients",
      "geography": "Where they operate",
      "volume": "Transactions/tasks/hires per period"
    }},
    "teamStructure": "Description of the team — roles, headcount, reporting structure",
    "toolsAndSystems": ["List of all tools, software, apps, systems they use"],
    "keyClients": ["Notable client names if mentioned"]
  }},
  "processMap": {{
    "processName": "Name of the core process",
    "steps": [
      {{
        "stepNumber": 1,
        "name": "Step name",
        "description": "Detailed description of what happens",
        "owner": "Who does this — role/team name",
        "teamSize": "How many people are involved (if known)",
        "toolsUsed": ["Tools used in this step"],
        "estimatedTime": "How long this step takes",
        "volume": "How often / how many (if known)",
        "painLevel": "low|medium|high",
        "painDescription": "What makes this painful (if applicable)",
        "automationPotential": "low|medium|high",
        "automationIdea": "How AI could automate or assist this step",
        "confidence": "high|medium|low",
        "notes": "Additional context"
      }}
    ]
  }},
  "keyInsights": ["Important patterns, observations, or strategic insights"],
  "topAutomationTargets": [
    {{
      "target": "What to automate",
      "currentCost": "Time/money/people currently spent",
      "automationApproach": "How AI could handle this",
      "expectedImpact": "What improvement to expect",
      "priority": "high|medium|low"
    }}
  ],
  "knowledgeGaps": ["Specific questions or topics where information is still missing"],
  "communicationChannels": ["How this BU communicates with workers"],
  "complianceNotes": "Any regulatory, legal, or compliance requirements mentioned"
}}

IMPORTANT RULES:
- If the previous process map exists and feedback was provided, incorporate ALL feedback
- Set confidence to 'low' for any step where you are inferring rather than directly told
- Be specific — avoid generic descriptions
- If information was NOT provided for a field, set it to null rather than guessing
- The knowledgeGaps section is critical — be thorough about what we still need to learn"""


@retry(
    stop=stop_after_attempt(3),
    wait=wait_exponential(multiplier=1, min=2, max=60),
    retry=retry_if_exception(is_rate_limit_error),
    reraise=True
)
def generate_process_map(vertical_id, user_id):
    vertical = Vertical.query.get(vertical_id)
    if not vertical:
        return None

    messages = Message.query.filter_by(vertical_id=vertical_id).order_by(Message.created_at).all()
    documents = Document.query.filter_by(vertical_id=vertical_id, processing_status='done').all()
    notes = Note.query.filter_by(vertical_id=vertical_id).all()
    prev_map = ProcessMap.query.filter_by(vertical_id=vertical_id).order_by(ProcessMap.version.desc()).first()

    context_parts = []

    if messages:
        context_parts.append("--- CONVERSATION TRANSCRIPT ---")
        for msg in messages:
            user = msg.user
            name = user.display_name if user and msg.role == 'user' else 'AI Analyst'
            context_parts.append(f"{name} ({msg.role}): {msg.content}")

    if documents:
        context_parts.append("\n--- UPLOADED DOCUMENTS (AI-extracted content) ---")
        for doc in documents:
            context_parts.append(f"Document: {doc.filename} (Type: {doc.doc_type})")
            if doc.extracted_content:
                context_parts.append(doc.extracted_content)

    if notes:
        context_parts.append("\n--- NOTES ---")
        for note in notes:
            context_parts.append(f"[{note.category}] {note.content}")

    if prev_map:
        context_parts.append("\n--- PREVIOUS PROCESS MAP ---")
        context_parts.append(prev_map.map_data)

        feedback_entries = ProcessMapFeedback.query.filter_by(process_map_id=prev_map.id).all()
        if feedback_entries:
            context_parts.append("\n--- USER FEEDBACK ON PREVIOUS MAP ---")
            for fb in feedback_entries:
                step_info = f"Step {fb.step_number}" if fb.step_number else "General"
                context_parts.append(f"{step_info} [{fb.feedback_type}]: {fb.content}")

    full_context = "\n".join(context_parts)

    system_prompt = get_process_map_prompt(vertical)

    response = client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=8192,
        system=system_prompt,
        messages=[{
            "role": "user",
            "content": f"CONTEXT PROVIDED:\n\n{full_context}\n\nPlease produce the comprehensive analysis."
        }]
    )

    result_text = response.content[0].text

    new_version = (prev_map.version + 1) if prev_map else 1
    source_summary = f"Generated from {len(messages)} messages, {len(documents)} documents, {len(notes)} notes"
    if prev_map:
        feedback_count = ProcessMapFeedback.query.filter_by(process_map_id=prev_map.id).count()
        source_summary += f", incorporating {feedback_count} feedback items from v{prev_map.version}"

    process_map = ProcessMap(
        vertical_id=vertical_id,
        generated_by=user_id,
        version=new_version,
        map_data=result_text,
        source_summary=source_summary
    )
    db.session.add(process_map)
    db.session.commit()

    return process_map


def compute_context_hash(vertical_id):
    messages = Message.query.filter_by(vertical_id=vertical_id).count()
    documents = Document.query.filter_by(vertical_id=vertical_id, processing_status='done').count()
    notes = Note.query.filter_by(vertical_id=vertical_id).count()
    feedback = IntelligenceFeedback.query.filter_by(vertical_id=vertical_id).count()
    latest_msg = Message.query.filter_by(vertical_id=vertical_id).order_by(Message.created_at.desc()).first()
    latest_ts = latest_msg.created_at.isoformat() if latest_msg else ""
    raw = f"{messages}:{documents}:{notes}:{feedback}:{latest_ts}"
    return hashlib.md5(raw.encode()).hexdigest()


def _flatten_extracted_json(ec):
    """Convert extracted_content JSON string to readable text."""
    try:
        parsed = json.loads(ec)
        if not isinstance(parsed, dict):
            return str(ec)
        readable_parts = []
        for key, val in parsed.items():
            if val is None or val == '' or val == []:
                continue
            if isinstance(val, list):
                items = []
                for v in val:
                    if isinstance(v, dict):
                        items.append("; ".join(f"{k}: {vv}" for k, vv in v.items() if vv))
                    else:
                        items.append(str(v))
                readable_parts.append(f"{key}:\n  " + "\n  ".join(items))
            elif isinstance(val, dict):
                sub_parts = []
                for k, v in val.items():
                    if v is None or v == '' or v == []:
                        continue
                    if isinstance(v, list):
                        sub_parts.append(f"  {k}: {', '.join(str(x) for x in v)}")
                    elif isinstance(v, dict):
                        nested = "; ".join(f"{nk}: {nv}" for nk, nv in v.items() if nv)
                        sub_parts.append(f"  {k}: {nested}")
                    else:
                        sub_parts.append(f"  {k}: {v}")
                if sub_parts:
                    readable_parts.append(f"{key}:\n" + "\n".join(sub_parts))
            else:
                readable_parts.append(f"{key}: {val}")
        return "\n".join(readable_parts)
    except (json.JSONDecodeError, TypeError):
        return ec


def gather_all_context(vertical_id):
    """Assemble all context for intelligence generation.
    Uses full_text as the primary document source for maximum detail,
    with extracted_content as supplementary structured analysis.
    """
    messages = Message.query.filter_by(vertical_id=vertical_id).order_by(Message.created_at).all()
    documents = Document.query.filter_by(vertical_id=vertical_id, processing_status='done').all()
    notes = Note.query.filter_by(vertical_id=vertical_id).all()

    context_parts = []

    if messages:
        context_parts.append("--- CONVERSATION TRANSCRIPT ---")
        for msg in messages:
            user = msg.user
            name = user.display_name if user and msg.role == 'user' else 'AI Analyst'
            context_parts.append(f"{name} ({msg.role}): {msg.content}")

    if documents:
        context_parts.append("\n--- UPLOADED DOCUMENTS ---")
        for doc in documents:
            context_parts.append(f"\n=== Document: {doc.filename} (Type: {doc.doc_type}) ===")
            if doc.user_description:
                context_parts.append(f"User description: {doc.user_description}")

            if doc.full_text:
                context_parts.append(f"\n[Full Document Text]\n{doc.full_text}")

            if doc.extracted_content:
                context_parts.append(f"\n[AI-Extracted Analysis]\n{_flatten_extracted_json(doc.extracted_content)}")

    if notes:
        context_parts.append("\n--- NOTES ---")
        for note in notes:
            context_parts.append(f"[{note.category}] {note.content}")

    return "\n".join(context_parts), len(messages), len(documents), len(notes)


INTELLIGENCE_SECTION_ORDER = [
    'businessProfile',
    'businessModelCanvas',
    'serviceBlueprint',
    'processMap',
    'teamStructure',
    'toolsInventory',
    'painPoints',
    'knowledgeGaps',
    'contextCoverage',
    'automationReadiness',
]

SECTION_LABELS = {
    'businessProfile': 'Business Profile',
    'businessModelCanvas': 'Business Model Canvas',
    'serviceBlueprint': 'Service Blueprint',
    'processMap': 'Process Map',
    'teamStructure': 'Team Structure',
    'toolsInventory': 'Tools & Systems',
    'painPoints': 'Pain Points',
    'knowledgeGaps': 'Knowledge Gaps',
    'contextCoverage': 'Context Coverage',
    'automationReadiness': 'Automation Readiness',
}

SECTION_MAX_TOKENS = {
    'businessProfile': 1500,
    'businessModelCanvas': 2500,
    'serviceBlueprint': 3000,
    'processMap': 2500,
    'teamStructure': 1500,
    'toolsInventory': 1200,
    'painPoints': 1500,
    'knowledgeGaps': 1200,
    'contextCoverage': 1000,
    'automationReadiness': 2500,
}

# Sections that need other sections' outputs as input
SYNTHESIS_SECTIONS = {'knowledgeGaps', 'automationReadiness'}

# contextCoverage is computed programmatically, no Claude call
COMPUTED_SECTIONS = {'contextCoverage'}

_COST_ESTIMATION_INSTRUCTIONS = """COST ESTIMATION:
For every ops activity where you know or can reasonably estimate headcount and time spent, calculate a monthly cost estimate.
Use these assumptions:
- India ops team: ~₹25,000/month per person (₹142/hour for 176 hours/month)
- Malaysia ops team: ~RM3,500/month per person
- Indonesia ops team: ~Rp5,000,000/month per person
- Multiply: (people × hours_per_day × 22 working_days) / 176 × monthly_salary

Always show your math in costBasis and set confidence based on evidence."""


def _incremental_preamble(previous_data):
    """Build the incremental-update preamble if a previous version exists."""
    if not previous_data:
        return ""
    prev_json = json.dumps(previous_data, indent=2, default=str)
    return f"""PREVIOUS ANALYSIS (your prior version — refine, do not discard):
{prev_json}

INSTRUCTIONS FOR UPDATE:
- Preserve and refine all existing insights that are still supported by the context.
- Add new insights from any new context that was not available before.
- If new information contradicts your previous analysis, update and briefly explain why.
- Do NOT reduce detail or lose information from the previous version unless it was wrong.
"""


def get_section_prompt(section_key, vertical, previous_data=None, other_sections=None):
    """Return a focused system prompt for generating a single intelligence section."""
    base = f"You are a senior process analyst at BetterPlace Group's AI Labs. Analyze all available context for {vertical.name} ({vertical.geography} -- {vertical.type}).\n\n{vertical.seed_context}\n\n"
    incremental = _incremental_preamble(previous_data)

    if section_key == 'businessProfile':
        return base + incremental + """Produce the BUSINESS PROFILE section. Return ONLY valid JSON (no markdown, no backticks):

{
  "whatTheyDo": "2-3 sentence description or null",
  "businessModel": "How they make money or null",
  "geography": "Where they operate or null",
  "scale": {
    "workerNetwork": "Size or null",
    "clientCount": "Number or null",
    "monthlyVolume": "Tasks/hires/shifts per month or null"
  },
  "keyClients": ["Client names"] or null,
  "teamSize": "Headcount and structure or null",
  "primaryLanguages": ["Languages"] or null,
  "communicationChannels": ["Channels"] or null
}

Be specific — use actual names, numbers, and details from the context. Set null for anything not mentioned."""

    elif section_key == 'businessModelCanvas':
        return base + incremental + """Produce the BUSINESS MODEL CANVAS section. Return ONLY valid JSON (no markdown, no backticks):

{
  "valueProposition": {
    "toClients": "What clients pay for — the core promise, or null",
    "toWorkers": "What workers/candidates/end users get, or null"
  },
  "revenueModel": {
    "pricingMechanism": "How pricing works or null",
    "averageTicketSize": "Typical deal size or null",
    "paymentTerms": "When clients pay or null",
    "marginStructure": "Where margin comes from and where it leaks or null"
  },
  "keyActivities": [
    {
      "activity": "Activity name",
      "category": "supply_acquisition|demand_fulfillment|operations|compliance|support",
      "peopleInvolved": "How many people and which roles, or null",
      "timePerWeek": "Estimated hours per week or null",
      "costIntensity": "low|medium|high",
      "automationReady": "low|medium|high",
      "whyItMatters": "Why this activity is critical"
    }
  ],
  "keyResources": {
    "people": "Team description or null",
    "technology": "Systems, apps, platforms or null",
    "data": "What unique data they have that could power AI or null",
    "relationships": "Key client or partner relationships or null"
  },
  "costStructure": {
    "biggestCostDrivers": ["Top 3-4 cost items"],
    "fixedVsVariable": "Mostly fixed or variable or null",
    "unitEconomics": "Revenue minus cost per unit, if known, or null"
  },
  "competitivePosition": {
    "competitors": ["Known competitors"] or null,
    "defensibility": "What makes them hard to replace or null",
    "vulnerability": "Where they could lose to competition or disruption or null"
  }
}

Analyze this as a BUSINESS, not just a process. Map how value flows through the business. For each key activity, estimate cost intensity (how much ops budget it consumes) and automation readiness (how suitable for AI). If this is a software business, adapt: clients = enterprise buyers, workers = end users, key activities = product development + sales + implementation."""

    elif section_key == 'serviceBlueprint':
        return base + incremental + f"""Produce the SERVICE BLUEPRINT section. Return ONLY valid JSON (no markdown, no backticks):

{{
  "processName": "Name of the core journey",
  "costSummary": {{
    "estimatedTotalMonthlyCost": "Approx total monthly ops cost string or null",
    "highestCostStages": ["Stage name — cost string"],
    "potentialSavings": "Estimated savings range from automation or null"
  }},
  "stages": [
    {{
      "stageName": "Stage name",
      "customerJourney": {{
        "action": "What the client/customer experiences",
        "goal": "Why they do it or null"
      }},
      "workerJourney": {{
        "action": "What the worker/candidate/end user does",
        "friction": "Where they struggle or wait, or null"
      }},
      "opsTeamWork": {{
        "action": "What the ops team does",
        "owner": "Role/team",
        "teamSize": "Number of people or null",
        "hoursPerDay": "Hours spent per day per person or null",
        "toolsUsed": ["Tools"],
        "isManual": true,
        "painPoint": "What's frustrating or null",
        "costEstimate": {{
          "monthlyHours": "Total person-hours per month or null",
          "estimatedMonthlyCost": "Formatted monthly cost or null",
          "costBasis": "Math behind the estimate or null",
          "confidence": "high|medium|low"
        }}
      }},
      "automationOpportunity": {{
        "readiness": "low|medium|high",
        "idea": "How AI could help or null",
        "expectedImpact": "Expected impact or null"
      }},
      "validationQuestions": ["Specific questions to verify this stage"]
    }}
  ]
}}

{_COST_ESTIMATION_INSTRUCTIONS}"""

    elif section_key == 'processMap':
        return base + incremental + """Produce the PROCESS MAP section. Return ONLY valid JSON (no markdown, no backticks):

{
  "processName": "Name of the core process",
  "steps": [
    {
      "stepNumber": 1,
      "name": "Step name",
      "description": "Detailed description",
      "owner": "Who does this",
      "teamSize": "How many people involved or null",
      "toolsUsed": ["Tools used"],
      "estimatedTime": "How long or null",
      "volume": "How often or null",
      "painLevel": "low|medium|high",
      "painDescription": "What makes this painful or null",
      "automationPotential": "low|medium|high",
      "automationIdea": "How AI could help or null",
      "confidence": "high|medium|low",
      "notes": "Additional context or null"
    }
  ]
}

Map the end-to-end operational process. Be specific with names, volumes, and time estimates from the context. Set confidence to 'low' for any step you are inferring rather than being directly told about."""

    elif section_key == 'teamStructure':
        return base + incremental + """Produce the TEAM STRUCTURE section. Return ONLY valid JSON (no markdown, no backticks):

[
  {
    "role": "Role name",
    "headcount": "Number or null",
    "responsibilities": "Key responsibilities",
    "processSteps": ["Which process steps they own"]
  }
]

Extract every role mentioned. Preserve exact headcounts and team sizes. If reporting lines are mentioned, note them in responsibilities."""

    elif section_key == 'toolsInventory':
        return base + incremental + """Produce the TOOLS & SYSTEMS INVENTORY section. Return ONLY valid JSON (no markdown, no backticks):

[
  {
    "name": "Tool name",
    "type": "Tool type (CRM, spreadsheet, mobile app, communication, etc.)",
    "usedIn": ["Process steps where used"],
    "usedBy": "Role/team that uses it"
  }
]

Include every tool, app, platform, spreadsheet, and communication channel mentioned. Include informal tools like WhatsApp groups or shared Google Sheets."""

    elif section_key == 'painPoints':
        return base + incremental + """Produce the PAIN POINTS section. Return ONLY valid JSON (no markdown, no backticks):

[
  {
    "severity": "high|medium|low",
    "title": "Pain point title",
    "currentEffort": "Time/money/people spent or null",
    "affectedProcess": "Which process step or null",
    "automationIdea": "How AI could help",
    "expectedImpact": "What improvement to expect"
  }
]

Only include pain points actually mentioned or strongly implied in the context. Rank by severity. Include cost impact wherever possible."""

    elif section_key == 'knowledgeGaps':
        other_summary = ""
        if other_sections:
            other_summary = "\n\nSECTIONS ALREADY ANALYZED:\n"
            for k, v in other_sections.items():
                if v:
                    other_summary += f"\n--- {SECTION_LABELS.get(k, k)} ---\n{json.dumps(v, indent=1, default=str)[:1500]}\n"

        return base + incremental + f"""Produce the KNOWLEDGE GAPS section. Return ONLY valid JSON (no markdown, no backticks):

["Specific question about what is still missing"]

Generate 5-15 specific, actionable questions about information that is MISSING and would be needed to design AI automation agents. Consider what has already been covered and what remains unknown.
{other_summary}

Focus on gaps that would block automation design: missing process details, unknown volumes, unclear decision criteria, undocumented exceptions, missing cost data, unknown integration points."""

    elif section_key == 'automationReadiness':
        other_summary = ""
        if other_sections:
            other_summary = "\n\nALL SECTIONS ANALYZED SO FAR:\n"
            for k, v in other_sections.items():
                if v:
                    other_summary += f"\n--- {SECTION_LABELS.get(k, k)} ---\n{json.dumps(v, indent=1, default=str)[:2000]}\n"

        return base + incremental + f"""Produce the AUTOMATION READINESS SCORECARD. Return ONLY valid JSON (no markdown, no backticks):

{{
  "overallScore": 72,
  "contextCompleteness": {{
    "score": 65,
    "detail": "Coverage summary"
  }},
  "processClarity": {{
    "score": 80,
    "detail": "Process clarity summary"
  }},
  "dataAvailability": {{
    "score": 60,
    "detail": "Data availability summary"
  }},
  "teamReadiness": {{
    "score": 75,
    "detail": "Team readiness summary"
  }},
  "topAutomationCandidates": [
    {{
      "process": "What to automate",
      "currentMonthlyCost": "Current cost string or null",
      "automationType": "Type of AI agent/system",
      "estimatedSavings": "Savings estimate or null",
      "prerequisite": "What we need before building",
      "timeToImplement": "Estimated time",
      "priority": 1
    }}
  ],
  "blockers": ["What is preventing automation today"],
  "recommendedNextSteps": ["Specific next actions to move forward"]
}}

Rate on 4 dimensions (0-100): Context Completeness, Process Clarity, Data Availability, Team Readiness.
Overall = weighted average (context 30%, process 30%, data 25%, readiness 15%).
Rank top automation candidates by (monthly cost x feasibility). Be realistic about blockers.
{other_summary}"""

    return base + "Return ONLY valid JSON."


def _compute_context_coverage_section(vertical_id):
    """Programmatically compute context coverage — no Claude call needed."""
    coverage_data = compute_discovery_coverage(vertical_id)
    if not coverage_data:
        return None
    total = sum(c['percentage'] for c in coverage_data)
    overall = int(total / len(coverage_data)) if coverage_data else 0
    return {
        "overall": overall,
        "topics": [
            {
                "name": c['phase'],
                "percentage": c['percentage'],
                "captured": c['percentage'] > 25,
                "status": c['status'],
                "detail": c.get('detail', '')
            }
            for c in coverage_data
        ]
    }


def _enrich_context_for_section(full_context, vertical_id, section_key):
    """Append feedback and validated facts relevant to a section."""
    enriched = full_context

    feedback_entries = IntelligenceFeedback.query.filter_by(vertical_id=vertical_id).all()
    section_feedback = [fb for fb in feedback_entries if fb.section == section_key or fb.section == 'general']
    if section_feedback:
        enriched += "\n\n--- USER CORRECTIONS/FEEDBACK ---"
        for fb in section_feedback:
            enriched += f"\n[{fb.feedback_type}] {fb.field_path or 'general'}: "
            if fb.corrected_value:
                enriched += f"Corrected to: {fb.corrected_value}"
            if fb.comment:
                enriched += f" Comment: {fb.comment}"

    validated_facts = build_validated_facts(vertical_id)
    if validated_facts:
        enriched += "\n\n--- VALIDATED FACTS ---"
        for fact in validated_facts:
            enriched += f"\n- {fact}"

    return enriched


@retry(
    stop=stop_after_attempt(3),
    wait=wait_exponential(multiplier=1, min=2, max=60),
    retry=retry_if_exception(is_rate_limit_error),
    reraise=True
)
def generate_single_section(vertical, section_key, full_context, previous_data=None, other_sections=None):
    """Generate a single intelligence section via a focused Claude call."""
    system_prompt = get_section_prompt(section_key, vertical, previous_data, other_sections)
    context = _enrich_context_for_section(full_context, vertical.id, section_key)

    max_context = 150000
    if len(context) > max_context:
        context = context[:max_context] + "\n\n[Context truncated due to size]"

    response = client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=SECTION_MAX_TOKENS.get(section_key, 2000),
        system=system_prompt,
        messages=[{
            "role": "user",
            "content": f"CONTEXT:\n\n{context}\n\nProduce the {SECTION_LABELS.get(section_key, section_key)} analysis."
        }]
    )

    result_text = clean_json_response(response.content[0].text)
    try:
        return json.loads(result_text)
    except (json.JSONDecodeError, TypeError):
        print(f"Failed to parse {section_key} section JSON, returning raw text")
        return result_text


def _upsert_section(vertical_id, section_key, section_data, context_hash, user_id):
    """Insert or update an IntelligenceSection row."""
    existing = IntelligenceSection.query.filter_by(
        vertical_id=vertical_id, section_key=section_key
    ).first()

    data_str = json.dumps(section_data, default=str) if not isinstance(section_data, str) else section_data

    if existing:
        existing.section_data = data_str
        existing.context_hash = context_hash
        existing.version = (existing.version or 0) + 1
        existing.generated_at = datetime.utcnow()
        existing.generated_by = user_id
    else:
        existing = IntelligenceSection(
            vertical_id=vertical_id,
            section_key=section_key,
            section_data=data_str,
            context_hash=context_hash,
            version=1,
            generated_by=user_id,
        )
        db.session.add(existing)

    db.session.commit()
    return existing


def assemble_composite_intelligence(vertical_id, user_id, context_hash):
    """Merge all IntelligenceSection rows into a single VerticalIntelligence record."""
    sections = IntelligenceSection.query.filter_by(vertical_id=vertical_id).all()
    composite = {}
    for sec in sections:
        if sec.section_data:
            try:
                composite[sec.section_key] = json.loads(sec.section_data)
            except (json.JSONDecodeError, TypeError):
                composite[sec.section_key] = sec.section_data

    intel = VerticalIntelligence(
        vertical_id=vertical_id,
        intelligence_data=json.dumps(composite, default=str),
        context_hash=context_hash,
        generated_by=user_id,
    )
    db.session.add(intel)
    db.session.commit()
    return intel


def generate_intelligence_sections_background(app, vertical_id, user_id, force=False, sections=None):
    """Generate intelligence section-by-section in a background thread."""
    _intel_generation_status[vertical_id] = {
        "status": "generating",
        "error": None,
        "sections_total": len(sections or INTELLIGENCE_SECTION_ORDER),
        "sections_completed": 0,
        "current_section": None,
        "completed_sections": [],
        "started_at": datetime.utcnow().isoformat(),
    }

    with app.app_context():
        try:
            vertical = Vertical.query.get(vertical_id)
            if not vertical:
                _intel_generation_status[vertical_id] = {"status": "failed", "error": "Vertical not found"}
                return

            current_hash = compute_context_hash(vertical_id)
            full_context, msg_count, doc_count, note_count = gather_all_context(vertical_id)

            if not full_context.strip():
                _intel_generation_status[vertical_id] = {"status": "failed", "error": "No context available"}
                return

            sections_to_gen = sections or list(INTELLIGENCE_SECTION_ORDER)
            completed = {}

            # Load existing sections that we are NOT regenerating (for synthesis inputs)
            all_existing = {s.section_key: s for s in IntelligenceSection.query.filter_by(vertical_id=vertical_id).all()}
            for key in INTELLIGENCE_SECTION_ORDER:
                if key not in sections_to_gen and key in all_existing and all_existing[key].section_data:
                    try:
                        completed[key] = json.loads(all_existing[key].section_data)
                    except (json.JSONDecodeError, TypeError):
                        pass

            for section_key in sections_to_gen:
                _intel_generation_status[vertical_id]["current_section"] = section_key

                # contextCoverage is computed, not generated by Claude
                if section_key in COMPUTED_SECTIONS:
                    result = _compute_context_coverage_section(vertical_id)
                    if result:
                        _upsert_section(vertical_id, section_key, result, current_hash, user_id)
                        completed[section_key] = result
                    _intel_generation_status[vertical_id]["sections_completed"] += 1
                    _intel_generation_status[vertical_id]["completed_sections"].append(section_key)
                    continue

                # Load previous version for incremental updates
                prev_data = None
                if section_key in all_existing and all_existing[section_key].section_data:
                    try:
                        prev_data = json.loads(all_existing[section_key].section_data)
                    except (json.JSONDecodeError, TypeError):
                        pass

                # Skip if hash matches and not forced
                if not force and prev_data and section_key in all_existing:
                    if all_existing[section_key].context_hash == current_hash:
                        completed[section_key] = prev_data
                        _intel_generation_status[vertical_id]["sections_completed"] += 1
                        _intel_generation_status[vertical_id]["completed_sections"].append(section_key)
                        continue

                # Synthesis sections get other sections as input
                other = completed if section_key in SYNTHESIS_SECTIONS else None

                result = generate_single_section(vertical, section_key, full_context, prev_data, other)
                if result:
                    _upsert_section(vertical_id, section_key, result, current_hash, user_id)
                    completed[section_key] = result

                _intel_generation_status[vertical_id]["sections_completed"] += 1
                _intel_generation_status[vertical_id]["completed_sections"].append(section_key)

            # Assemble composite
            assemble_composite_intelligence(vertical_id, user_id, current_hash)
            _intel_generation_status[vertical_id]["status"] = "done"
            _intel_generation_status[vertical_id]["current_section"] = None

        except Exception as e:
            print(f"Intelligence generation error for {vertical_id}: {e}")
            _intel_generation_status[vertical_id] = {"status": "failed", "error": str(e)}


def start_intelligence_generation(app, vertical_id, user_id, force=True, sections=None):
    """Start async per-section intelligence generation — returns immediately."""
    thread = threading.Thread(
        target=generate_intelligence_sections_background,
        args=(app, vertical_id, user_id, force, sections)
    )
    thread.daemon = True
    thread.start()


# Incremental intelligence after new context arrives
_last_incremental_trigger = {}
INCREMENTAL_COOLDOWN_SECONDS = 60


def start_incremental_intelligence(app, vertical_id, user_id=None):
    """Trigger incremental intelligence update if prior intelligence exists and cooldown has passed."""
    import time
    now = time.time()
    last = _last_incremental_trigger.get(vertical_id, 0)
    if now - last < INCREMENTAL_COOLDOWN_SECONDS:
        return False

    status = _intel_generation_status.get(vertical_id, {})
    if status.get("status") == "generating":
        return False

    has_prior = IntelligenceSection.query.filter_by(vertical_id=vertical_id).first()
    if not has_prior:
        return False

    _last_incremental_trigger[vertical_id] = now
    start_intelligence_generation(app, vertical_id, user_id, force=False, sections=None)
    return True


@retry(
    stop=stop_after_attempt(3),
    wait=wait_exponential(multiplier=1, min=2, max=60),
    retry=retry_if_exception(is_rate_limit_error),
    reraise=True
)
def generate_cross_vertical_analysis():
    all_intelligence = []
    verticals = Vertical.query.all()

    for vertical in verticals:
        intel = VerticalIntelligence.query.filter_by(
            vertical_id=vertical.id
        ).order_by(VerticalIntelligence.generated_at.desc()).first()
        if intel and intel.intelligence_data:
            all_intelligence.append({
                "vertical": vertical.name,
                "vertical_id": vertical.id,
                "type": vertical.type,
                "geography": vertical.geography,
                "intelligence": intel.intelligence_data,
            })

    if len(all_intelligence) < 2:
        return None

    prompt = """You are analyzing intelligence data from multiple business units within BetterPlace Group. Your goal is to identify CROSS-VERTICAL PATTERNS that can inform shared automation strategies.

Return ONLY valid JSON:
{
  "commonPainPoints": [
    {
      "painPoint": "Shared pain point",
      "affectedVerticals": ["Vertical names"],
      "combinedMonthlyCost": "Combined cost string or null",
      "sharedAutomationOpportunity": "One agent or approach that could serve all affected verticals"
    }
  ],
  "sharedProcessPatterns": [
    {
      "pattern": "Process pattern name",
      "description": "How this pattern manifests across verticals",
      "verticals": ["Vertical names"],
      "commonSteps": ["Similar steps"],
      "differences": ["Key differences by geography or business model"],
      "reusableAgentDesign": "How one agent design could be reused"
    }
  ],
  "automationPriorityMatrix": [
    {
      "automationTarget": "What to automate",
      "verticals": ["Which verticals benefit"],
      "totalMonthlyCost": "Combined cost",
      "totalEstimatedSavings": "Combined savings",
      "buildOnceServeMany": true,
      "recommendedBeachhead": "Which vertical to start with and why",
      "priority": 1
    }
  ],
  "uniqueInsightsPerVertical": [
    {
      "vertical": "Vertical name",
      "uniqueAspect": "Unique trait",
      "implication": "What it means for automation strategy"
    }
  ],
  "overallRecommendation": "2-3 sentence summary of the optimal automation strategy across all verticals"
}

IMPORTANT:
- OkayGo and Troopers are the same business model in different geographies; maximize reuse opportunities.
- MyRobin is similar but with longer-term placements.
- AasaanJobs shares sourcing/screening patterns with staffing businesses.
- Background Verification is a cross-cutting function.
- goBetter is software and should likely be treated separately from services ops.

Highlight the highest-leverage build-once-serve-many opportunities."""

    context = "\n\n".join(
        f"=== {item['vertical']} ({item['geography']} — {item['type']}) ===\n{item['intelligence']}"
        for item in all_intelligence
    )

    response = client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=4096,
        system=prompt,
        messages=[{"role": "user", "content": context}]
    )

    return response.content[0].text


def generate_automation_brief(intelligence_data, export_data):
    intelligence = safe_json_loads(intelligence_data, default={})
    if not intelligence:
        return ""

    vertical = export_data.get("vertical", {})
    business_canvas = intelligence.get("businessModelCanvas") or {}
    service_blueprint = intelligence.get("serviceBlueprint") or {}
    automation = intelligence.get("automationReadiness") or {}
    gaps = intelligence.get("knowledgeGaps") or []
    validated_facts = export_data.get("validated_facts") or []

    lines = [
        f"# {vertical.get('name', 'Unknown Vertical')} — AI Automation Brief",
        f"Generated from Context Brain on {export_data.get('exported_at', '')}",
        "",
        "## Business Model Summary",
        f"- Client value proposition: {((business_canvas.get('valueProposition') or {}).get('toClients')) or 'Not yet captured'}",
        f"- Worker/user value proposition: {((business_canvas.get('valueProposition') or {}).get('toWorkers')) or 'Not yet captured'}",
        f"- Pricing mechanism: {((business_canvas.get('revenueModel') or {}).get('pricingMechanism')) or 'Not yet captured'}",
        f"- Margin structure: {((business_canvas.get('revenueModel') or {}).get('marginStructure')) or 'Not yet captured'}",
        "",
        "## Service Blueprint Summary",
    ]

    for stage in (service_blueprint.get("stages") or [])[:8]:
        ops = stage.get("opsTeamWork") or {}
        cost = (ops.get("costEstimate") or {}).get("estimatedMonthlyCost") or "Cost not estimated"
        lines.append(
            f"- {stage.get('stageName', 'Stage')}: {ops.get('action') or 'No ops action captured'} — {cost}"
        )

    lines.extend([
        "",
        "## Top Automation Targets (Ranked)",
    ])

    for idx, candidate in enumerate((automation.get("topAutomationCandidates") or [])[:5], start=1):
        lines.append(
            f"{idx}. {candidate.get('process', 'Unknown target')} — Cost: {candidate.get('currentMonthlyCost') or 'Unknown'} — Savings: {candidate.get('estimatedSavings') or 'Unknown'} — Time: {candidate.get('timeToImplement') or 'Unknown'}"
        )
        lines.append(f"   Prerequisites: {candidate.get('prerequisite') or 'None captured'}")
        lines.append("")

    lines.extend([
        f"## Automation Readiness Score: {automation.get('overallScore') or 0}/100",
        f"- Context: {((automation.get('contextCompleteness') or {}).get('score')) or 0}% | Process: {((automation.get('processClarity') or {}).get('score')) or 0}% | Data: {((automation.get('dataAvailability') or {}).get('score')) or 0}% | Team: {((automation.get('teamReadiness') or {}).get('score')) or 0}%",
        f"- Blockers: {', '.join(automation.get('blockers') or ['None captured'])}",
        "",
        "## Knowledge Gaps (Still Need)",
    ])

    if gaps:
        for gap in gaps:
            lines.append(f"- {gap.get('question') if isinstance(gap, dict) else gap}")
    else:
        lines.append("- None captured")

    lines.extend([
        "",
        "## Full Service Blueprint",
    ])

    for stage in service_blueprint.get("stages") or []:
        customer = stage.get("customerJourney") or {}
        worker = stage.get("workerJourney") or {}
        ops = stage.get("opsTeamWork") or {}
        lines.append(f"### {stage.get('stageName', 'Stage')}")
        lines.append(f"- Customer: {customer.get('action') or 'Not captured'}")
        lines.append(f"- Worker: {worker.get('action') or 'Not captured'}")
        lines.append(f"- Ops: {ops.get('action') or 'Not captured'}")
        lines.append(f"- Cost: {((ops.get('costEstimate') or {}).get('estimatedMonthlyCost')) or 'Not estimated'}")
        lines.append("")

    lines.extend([
        "## Raw Context Summary",
        f"- {len(export_data.get('conversation') or [])} chat messages captured",
        f"- {len(export_data.get('documents') or [])} documents processed",
        f"- {len(export_data.get('notes') or [])} notes added",
        f"- Key validated facts: {', '.join(validated_facts) if validated_facts else 'None captured'}",
        "",
        "## Recommended Agent Design Starting Points",
    ])

    for candidate in automation.get("topAutomationCandidates") or []:
        lines.append(
            f"- {candidate.get('process', 'Unknown target')}: {candidate.get('automationType') or 'No agent type captured'}"
        )

    return "\n".join(lines).strip() + "\n"
